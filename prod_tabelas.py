import os
import glob
from argparse import ArgumentParser
from datetime import datetime as dt

from IPython.display import Image, HTML
from bs4 import BeautifulSoup as BS
import pandas as pd
#from jinja2 import Environment, FileSystemLoader

import imgkit
import pdfkit
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH

imgkit_options = {'encoding': "UTF-8",'enable-local-file-access': None, "quality": 100}

# dicionario com os codigos dos escudos dos clubes no site da CBF (possivel mudar pra uma base local com escudos em png)
codtime = {"Athletico Paranaense":"4b3e1810-ed20-4d66-9f09-fa50d654aa26",
    "Athletico-PR":"4b3e1810-ed20-4d66-9f09-fa50d654aa26",
    "Atlético-MG":"e75ea650-c770-4d07-9403-a94c7cd44cd0",
    "Ceará":"35767326-ada8-453c-b18d-5f1aec90a765",
    "Fortaleza":"0a776ee7-565e-4988-bba0-09f153b54be5",
    "Atlético-GO":"4c8569a8-872e-4cf8-8bb2-2cdc2afcbaef",
    "Bahia":"e1827e57-2f78-4a81-9678-9011ab63a43d",
    "Botafogo":"a30426cf-4864-4a24-9cc9-8caf6b6d0881",
    "Coritiba":"aac40e6b-3a4c-478a-a93b-caf586a831a9",
    "Flamengo":"78061125-cb2a-49f5-a556-4eb1cda03877",
    "Vasco":"d2f54509-8176-4a24-aed8-55dde475de75",
    "Vasco da Gama":"d2f54509-8176-4a24-aed8-55dde475de75",
    "Grêmio":"6ffc058a-9bec-4ca5-968d-114e352e9e7a",
    "Goiás":"1d77fc2e-cdf4-4b82-af63-c48e5fbe99f6",
    "Fluminense":"32e8e52f-14f2-4306-998d-53fa8d205b71",
    "Bragantino":"0a87c7cf-cdbf-4d11-9e1c-741619b9f60c",
    "Internacional":"568dd980-6fa7-4212-b13a-39d0ab876148",
    "Sport":"275c970b-b403-49c7-8c65-63c927592911",
    "São Paulo":"9aa4a5ed-c26f-45bb-b13e-6a7196f05fb7",
    "Santos":"54917956-77d8-4436-898c-29a88c1c31ff",
    "Palmeiras":"dd44c5aa-329f-4a96-ab9c-d8de6861e299",
    "Corinthians":"e65dc940-6c2c-4d9c-b501-4dd23492cd22",
    "Chapecoense":"00002sc",
    "Athlético-PR":"4b3e1810-ed20-4d66-9f09-fa50d654aa26"}

def highlight_green(s, media_dict, std_dict):
    estilo_verde = 'color: #00B050 ; font-size: 160%'
    key = ''.join(reversed(''.join((reversed(s.name))).split('_', 1)[-1]))
    try:
        up_std = 10E-4 > - s + (media_dict.get(key, 0) + std_dict.get(key, 0))
        return [estilo_verde if v else '' for v in up_std]
    except:
        return ""

def highlight_red(s, media_dict, std_dict):
    estilo_vermelho = 'color: #C00000 ; font-size: 160%'
    key = ''.join(reversed(''.join((reversed(s.name))).split('_', 1)[-1]))
    try:
        down_std = 10E-4 < - s + (media_dict.get(key, 0) - std_dict.get(key, 0))
        return [estilo_vermelho if v else '' for v in down_std]
    except:
        return ""
############################
# coloquei essas duas funcoes aqui + as linhas 179 e 180
def highlight_red_gols(s, media_dict, std_dict):
    estilo_verde = 'color: #C00000 ; font-size: 160%'
    key = ''.join(reversed(''.join((reversed(s.name))).split('_', 1)[-1]))
    try:
        up_std = 10E-4 > - s + (media_dict.get(key, 0) + std_dict.get(key, 0))
        return [estilo_verde if v else '' for v in up_std]
    except:
        return ""

def highlight_green_gols(s, media_dict, std_dict):
    estilo_vermelho = 'color: #00B050 ; font-size: 160%'
    key = ''.join(reversed(''.join((reversed(s.name))).split('_', 1)[-1]))
    try:
        down_std = 10E-4 < - s + (media_dict.get(key, 0) - std_dict.get(key, 0))
        return [estilo_vermelho if v else '' for v in down_std]
    except:
        return ""    
############################    
def highlight_nmax(s, n=3):
    estilo_verde = 'color: #00B050 ; font-size: 160%'
    n_largest = s.nlargest(n).min()
    try:
        up = n_largest - s < 10E-4 
        return [estilo_verde if v else '' for v in up]
    except:
        return ""


def highlight_nmin(s, n=3):
    estilo_verde = 'color: #00B050 ; font-size: 160%'
    n_small = s.nsmallest(n).max()
    try:
        down = s - n_small < 10E-4 
        return [estilo_verde if v else '' for v in down]
    except:
        return ""


def field_formatter(key):
    """Configuração de formatadores para usar em pandas.Styler para colunas com nomes 
    únicos. Nome deve ser exatamente como vem junto ao dado de entrada.
    """
    varname = ''.join(reversed(''.join((reversed(key))).split('_', 1)[-1]))
    format_dict = {
        "SG":"{:.0%}",
        "SGced":"{:.0%}",
        "prob_vit":"{:.0%}",
        "prob_2+":"{:.0%}",
        "prob_SG":"{:.0%}",
        "Dentro da área":"{:.0f}",
        "Fora da área":"{:.0f}",
        "Penalti":"{:.0f}",
        "Gol contra":"{:.0f}",
        "Falta":"{:.0f}",
        "Cabeça":"{:.0f}",
        "GOL":"{:4.2f}",
        "LAT":"{:4.2f}",
        "MEI":"{:4.2f}",
        "ZAG":"{:4.2f}",
        "ATA":"{:4.2f}",
        }
    return format_dict.get(varname, "{}")
#funcao pra criar as tabelas

def criar_tabelas(arq_tabela, rodada, outfileName="", css_file="."):
    
    nome_tabela = os.path.basename(os.path.splitext(arq_tabela)[0])
    # leitura dos txt do cartola.py
    df = pd.read_csv(arq_tabela,
                    encoding='utf8', 
                    header = 0, 
                    index_col = 0,
                    sep=";")
                    


    # troca virgula por ponto, se necessario
    try:
        df = df.stack().str.replace(',','.').unstack()
    except:
        pass
    df = df.apply(pd.to_numeric, errors="ignore")
    medias = df.mean().to_dict()
    desvios = df.std().to_dict()

    #leitura das partidas da rodada
    partidas = pd.read_csv(os.path.join(DATA_DIR, 'partidas_#{}.txt'.format(rodada)),
                    encoding='utf8', 
                    header = 0, 
                    sep='\t') 
    
# desenha o df de acordo com o tipo de tabela
    
    if 'goleiros' in arq_tabela: 
    #ajustes na tabela (ordem alfabetica, index, e datatype)
        df= df.sort_values('CLUBE')
        df = df[['CLUBE','LOCAL','DD_MED','GS_MED','DD_CED','GP_ADV','F_TOT_ADV',"ADV"]]
        df.rename(columns={"CLUBE":"Time","LOCAL":"Mando", "DD_MED": "DDs", 'GS_MED': "Gols Sofridos", 'DD_CED':"DDs cedidas pelo ADV", 'GP_ADV': "Gols do ADV", 'F_TOT_ADV': "Finalizações do ADV"  },inplace=True)
        medias = df.mean().to_dict()
        desvios = df.std().to_dict()
    #colocando os escudos no lugar dos nomes e tornando-os index do dataframe
        escudos_adv = ["<img src=https://mcusercontent.com/ece121c367dd18e1b76860f4d/images/{}.png width = '40'>".format(codtime[i[1]["ADV"]])for i in df.iterrows()]
        df["ADV"] = escudos_adv
        escudos_time = ["<img src=https://mcusercontent.com/ece121c367dd18e1b76860f4d/images/{}.png width = '40'>".format(codtime[i[1]["Time"]])for i in df.iterrows()]
        df["Time"] = escudos_time
        df.index.name=None
        
    #formatacao condicional       
        html=df.style\
          .apply(highlight_green, media_dict = medias, std_dict = desvios, subset=['DDs','DDs cedidas pelo ADV','Finalizações do ADV'])\
          .apply(highlight_red, media_dict = medias, std_dict = desvios, subset=['DDs','DDs cedidas pelo ADV','Finalizações do ADV'])\
          .set_precision(2)\
          .set_table_attributes("class=tabela")\
          .apply(highlight_green_gols, media_dict = medias, std_dict = desvios, subset=['Gols Sofridos','Gols do ADV'])\
          .apply(highlight_red_gols, media_dict = medias, std_dict = desvios, subset=['Gols Sofridos','Gols do ADV'])
    else:
        
        df_man = df.reindex(partidas.MANDANTE).reset_index()
        df_man = df_man.rename(columns={'MANDANTE':'Man'})
        df_vis = df.reindex(partidas.VISITANTE).reset_index()
        df_vis = df_vis.rename(columns={'VISITANTE':'Vis'})
     
        ced_cruzada = pd.merge(df_man, right=df_vis, right_index=True, left_index=True, suffixes=("_Man","_Vis"))
        if 'desempenho' in arq_tabela:
            ordem = ["SG","GP","GS","SGced"]
        elif 'gols_por_forma' in arq_tabela:
            ordem = ['Dentro da área', "Cabeça", "Falta","Fora da área", "Penalti", "Gol contra"]
        elif 'scoreOdds' in arq_tabela:
            ordem = ["prob_vit", "prob_2+", "prob_SG"]
        else:
            ordem = ["GOL","LAT","ZAG","MEI","ATA"]
        
        ordem_final = []
        for t in ['Man','Vis']:
            temp = ["{}_{}".format(item, t) for item in ordem] + [t]
            if t =='Vis':
                temp = list(reversed(temp))
            ordem_final.extend(temp)

        ced_cruzada = ced_cruzada.reindex(ordem_final, axis=1)
        #ced_cruzada.columns = ced_cruzada.columns.str.split('_').str[0]
        df = ced_cruzada.copy()
    
    #colocando os escudos no lugar dos nomes e tornando-os index do dataframe
        escudos_adv = ["<img src=https://mcusercontent.com/ece121c367dd18e1b76860f4d/images/{}.png width = '60'>".format(codtime[i[1]["Vis"]])for i in df.iterrows()]
        df["Vis"] = escudos_adv
        escudos_time = ["<img src=https://mcusercontent.com/ece121c367dd18e1b76860f4d/images/{}.png width = '60'>".format(codtime[i[1]["Man"]])for i in df.iterrows()]
        df["Man"] = escudos_time
        df = df.fillna(0)

        
        # formatacao condicional 
        format_str = dict([(k, field_formatter(k)) for k in df.columns])
        
        # GS_posicao tem configuracao especial
        if 'GS_por_posicao' in arq_tabela:
            df = df.drop(["GOL_Man","GOL_Vis"], axis=1)
            for k in format_str.keys():
                if k not in ['Man', 'Vis']:
                    format_str[k] = "{:.0f}"

        html = df.style\
                 .apply(highlight_green, media_dict=medias, std_dict = desvios)\
                 .apply(highlight_red, media_dict=medias, std_dict = desvios)\
                 .set_precision(2)\
                 .set_table_attributes("class=tabela")\
                 .hide_index()\
                 .format(format_str)

    label_change = {"SG":"SG Obtidos", 
                    "GP":"Gols pro", 
                    "GS":"Gols Sofridos", 
                    "SGced":"SG Cedidos", 
                    "prob_vit":"Chance de vitória", 
                    "prob_2+":"Chance de 2 ou mais gols", 
                    "prob_SG":"Chance de SG",
                   }
    
    style_sheet = "file:///" +os.path.abspath(css_file)
    # edita o html pra ficar compativel com o CSS
    content = html.render()
    soup = BS(content, "lxml")
    if "goleiros" in arq_tabela:
        soup.find("th", {"class":"blank level0"}).insert(1,"Goleiros")
    else:
        for th in soup.find_all('th'):
            # reversed split for cases where column name contains '_'
            head = ''.join(reversed(''.join((reversed(th.string))).split('_', 1)[-1]))
            head = label_change.get(head, head)
            th.string = head
    new_tag = soup.new_tag("link",rel="stylesheet", href=style_sheet, type='text/css')
    soup.find("head").insert(1,new_tag)

    outname = outfileName or nome_tabela.format(rodada)
    # cria o arquivo e renderiza o html
    with open(os.path.join(path, outname + '.html'),'w',encoding='utf8') as f:
        f.write(str(soup))

    return df


def html_to_jpg(repopath, css, outpath=""):    
    outpath = outpath or repopath
    
    for arq in os.listdir(repopath):
        fname, ext = os.path.splitext(arq)
        if "html" in ext:    
            imgkit.from_file(os.path.join(repopath, arq), 
                            os.path.join(outpath, fname + '.jpg'),
                            css=os.path.abspath(css),
                            options=imgkit_options)


def gera_pdf(imgdir, rodada, outpath=os.getcwd(),**kwargs):
    """cria o arquivo pros socios
    """

    # TODO: Organizar arquivos em grupos de pdfs, config aqui dentro, texto descritivo vem como arg
    # TODO: Colocar para receber exemplos de dados para escrever nos textos. Receber texto exemplo pronto de fora
    # TODO: colocar exemplo como caption da tabela com styler do pandas

    textos ={"cedidos_#{rodada:d}": "PONTOS CEDIDOS: Média dos pontos cedidos pelas equipes aos adversários no "
                                    "Cartola em {ano}, considerando o mando de campo. Em destaque, as equipes com"
                                    "desvio acima padrão.",
            "cedidos(semMando)_#{rodada:d}": "PONTOS CEDIDOS SEM MANDO: Média dos pontos cedidos pelas equipes aos adversários no "
                                    "Cartola em {ano}, considerando todos seus jogos SEM distinção de mando de campo. Em destaque, "
                                    "as equipes com maiores distância para média do campeoanto.",
            "liqCed_#{rodada:d}": "PONTOS CEDIDOS LÍQUIDOS (EXCLUSIVO): Média dos pontos cedidos pelas equipes "
                                  "aos adversários no Cartola em {ano}, tirando da conta as pontuações de gols, "
                                  "assistências e SG. A tabela não considera o mando de campo. Em destaque, as "
                                  "equipes com desvio acima padrão.",
            "ptosPos_#{rodada:d}": "PONTOS GANHOS (EXCLUSIVO): Média de pontos ganhos pelas equipes, separados por "
                                   "posição, no Cartola FC. A tabela não faz a distinção de mando de campo.",
            "apostas_#{rodada:d}": "CASAS DE APOSTAS (PRIVADO): Análise da probabilidade de vitória, SG e gols dos "
                                   "confrontos nas principais casas de apostas.",
            "rbCed_#{rodada:d}": "DSs e DDs CEDIDAS: Média dos desarmes e defesas difíceis (goleiros) cedidos pelas "
                                 "equipes aos adversários no Cartola em {ano}, sem considerar o mando de campo. Em destaque, "
                                 "as equipes com desvio acima do padrão.",
            "desempenhoClubes_#{rodada:d}": "DESEMPENHO DOS TIMES: Desempenho dos times no Cartola FC em {ano}, sem levar em "
                                            "consideração o mando de campo. As informações são a média de: SG, Gols Pró, Gols"
                                            " Sofridos e SG Cedidos.",
            'goleiros_#{rodada:d}': "GOLEIROS (EXCLUSIVO): Médias por jogo dos goleiros no Cartola FC em {ano}. Todos os valores"
                                    " em média por jogo. A tabela conta com distinção de mando de campo.",
            'GS_Forma_#{rodada:d}': "GOLS SOFRIDOS POR FORMA: Tabela com os tipos de gols sofridos por equipe do Campeonato "
                                    "Brasileiro {ano}. A tabela conta todos os jogos (casa e fora).",
            'GS_posicao_#{rodada:d}': "GOLS SOFRIDOS POR POSIÇÃO (EXCLUSIVO): Tabela com os gols sofridos por equipe do Campeonato "
                                      "Brasileiro {ano}, separados pela posição do jogador adversário que marcou o gol. A tabela conta"
                                      " todos os jogos (casa e fora).",
    }
    
    doc=Document(".\\pytabelas\\socios.docx")
    p = doc.add_paragraph()
    r = p.add_run()
    font = r.font
    font.name="Industry-Black"
    font.size = Pt(36)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r.add_text("ARQUIVO DOS SÓCIOS RODADA " +str(rodada))

    for prod, decr in textos.items():
        r.add_break(WD_BREAK.PAGE)
        p = doc.add_paragraph()
        r = p.add_run()
        font = r.font
        font.name="Industry-Black"
        font.size = Pt(24)
        r.add_text(decr.format(rodada=rodada, **kwargs))
        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(imgdir, prod.format(rodada=rodada, **kwargs)+".jpg"), width= Inches(6.1))
        
        
    doc.save(os.path.join(outpath, 'socios #{}'.format(rodada)+'.docx'))

parser = ArgumentParser('prod_tabelas', 
                        description="script para producao de tabelas formatadas e prontas"
                                    " para publicacao em redes do estats cartola."
                       )
parser.add_argument("--rodada", "-r", type=int, default=1, dest='rodada')
parser.add_argument("--ano", "-a", type=int, default=dt.now().year, dest='ano')
#parser.add_argument("--tabela", choices=["cedidos"])

#env = Environment(loader = FileSystemLoader(repo),)
#repo = '.\\templates'

#lista das tabelas a serem feitas
tab_cedidos = ["rbCed_#{}.txt","liqCed_#{}.txt",'cedidos_#{}.csv','ptosPos_#{}.csv','desempenhoClubes_#{}.txt',
               "ult10Jogs_liqCed_#{}.txt","ult10Jogs_rbCed_#{}.txt","ult10Jogs_PtosPos_#{}.txt",
               'desempenhoClubesUlt10Jogos_#{}.txt','goleiros_#{}.txt',"cedidos(semMando)_#{}.csv",
               "ult10Jogs_PtsCed_#{}.txt", "ult10_PtosCedidos(semMando)_#{}.csv"]

tab_scouts = ['GS_por_posicao.csv', "gols_por_forma_pivoted.csv"]
tab_scouts_labels = dict(zip(tab_scouts, ["GS_posicao_#{}", 'GS_Forma_#{}']))

if __name__ == "__main__":
    args = parser.parse_args()
    rodada = args.rodada  # pegar esse valor do API do Cartola
    ano = args.ano
    DATA_DIR = ".\\data\\{:4d}\\Rodada{:d}".format(ano, rodada)
    SCOUTS_DIR = ".\\data\\{:4d}\\scouts".format(ano)
    css_file = "./pytabelas/css/style.css"

    # cria a pasta pra rodada
    path = '.\\#Conteúdo Ativo (Posts) {}\\Rodada{}\\'.format(ano, rodada) 
    if not os.path.exists(path):
        os.makedirs(path)

    for tab in tab_cedidos:
        try:
            df = criar_tabelas(os.path.join(DATA_DIR, tab.format(rodada)), rodada, css_file=css_file)
        except FileNotFoundError:
            print('A tabela '+tab.format(rodada)+ ' nao existe pra essa rodada')

    for tab in tab_scouts:
        try:
            outfile = tab_scouts_labels.get(tab, os.path.splitext(tab)[0])
            df = criar_tabelas(os.path.join(SCOUTS_DIR, tab), rodada, outfile.format(rodada), css_file=css_file)
        except FileNotFoundError:
            print('A tabela '+tab+ ' nao existe pra essa rodada')
    
    files_apostas = glob.glob('.\\data\\scoreOdds\\*.csv')
    mais_recente = max(files_apostas, key=os.path.getctime)
    criar_tabelas(mais_recente, rodada, "apostas_#{}".format(rodada), css_file=css_file)

    html_to_jpg(path, css=css_file)

    gera_pdf(path, rodada, path, ano=ano)

